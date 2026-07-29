"""
parser/docx_parser.py
-----------------------
Extracts raw text from DOCX resumes using python-docx.

Handles paragraphs and table cells (many resume templates put
skills or contact info in tables), preserving reading order as
closely as python-docx's API allows.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from utils.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_docx(path: str | Path) -> tuple[str, list[str]]:
    """
    Extract raw text from a DOCX file.

    Returns a tuple of (text, warnings).
    """
    path = Path(path)
    warnings: list[str] = []
    text_parts: list[str] = []

    try:
        document = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        logger.error("Failed to open DOCX {}: {}", path.name, exc)
        warnings.append(
            "This file could not be opened as a Word document. It may be "
            "corrupted or saved in an unsupported format (e.g. .doc "
            "instead of .docx)."
        )
        return "", warnings

    # Paragraphs (body text) in document order
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Tables — many resume templates put contact info / skills in a table
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)

    text = "\n".join(text_parts)

    if not text.strip():
        warnings.append("No extractable text was found in this document.")

    return text, warnings
